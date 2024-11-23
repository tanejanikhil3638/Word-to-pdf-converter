from flask import Flask, request, jsonify, render_template, send_file, redirect, url_for
from werkzeug.utils import secure_filename
import os
from docx import Document
from fpdf import FPDF
from PyPDF2 import PdfReader, PdfWriter

app = Flask(__name__)

# Define folders
UPLOAD_FOLDER = os.path.join(os.getcwd(), 'uploads')
CONVERTED_FOLDER = os.path.join(os.getcwd(), 'converted')

# Create necessary directories
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(CONVERTED_FOLDER, exist_ok=True)

# Configure Flask to serve static and template files
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['CONVERTED_FOLDER'] = CONVERTED_FOLDER


@app.route('/')
def home():
    return render_template('index.html')


@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400

    file = request.files['file']
    password = request.form.get('password', None)  # Get the password (optional)

    # Validate file type
    if not file.filename.endswith('.docx'):
        return jsonify({'error': 'Invalid file type. Please upload a .docx file.'}), 400

    filename = secure_filename(file.filename)
    file_path = os.path.join(UPLOAD_FOLDER, filename)

    # Save the uploaded file
    file.save(file_path)

    try:
        # Extract metadata and convert to PDF
        doc = Document(file_path)
        metadata = {
            'paragraph_count': len(doc.paragraphs),
            'word_count': sum(len(paragraph.text.split()) for paragraph in doc.paragraphs),
        }

        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for paragraph in doc.paragraphs:
            pdf.multi_cell(0, 10, paragraph.text)

        # Save PDF to file
        converted_file_path = os.path.join(CONVERTED_FOLDER, filename.replace('.docx', '.pdf'))
        pdf.output(converted_file_path)

        # If password provided, apply password to the PDF
        if password:
            protected_pdf_path = os.path.join(CONVERTED_FOLDER, filename.replace('.docx', '_protected.pdf'))

            with open(converted_file_path, "rb") as f:
                reader = PdfReader(f)
                writer = PdfWriter()

                for page in reader.pages:
                    writer.add_page(page)

                # Encrypt the PDF with the password
                writer.encrypt(password)

                # Write the password-protected PDF to a new file
                with open(protected_pdf_path, "wb") as output_file:
                    writer.write(output_file)

            # Return download URL for the password-protected PDF
            return jsonify({
                'message': 'File uploaded, converted, and password-protected successfully!',
                'metadata': metadata,
                'download_url': url_for('download_file', filename=filename.replace('.docx', '_protected.pdf'), _external=True)
            })
        else:
            # If no password, just return the normal PDF URL
            return jsonify({
                'message': 'File uploaded and converted successfully!',
                'metadata': metadata,
                'download_url': url_for('download_file', filename=filename.replace('.docx', '.pdf'), _external=True)
            })

    except Exception as e:
        return jsonify({'error': f'Error during processing: {str(e)}'}), 500


@app.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    try:
        file_path = os.path.join(CONVERTED_FOLDER, filename)
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True)
        else:
            return jsonify({'error': 'File not found'}), 404
    except Exception as e:
        return jsonify({'error': f'Error during download: {str(e)}'}), 500


if __name__ == '__main__':
    app.run(debug=True)
